"""Synthèse direction — 1-pager Go / No-go pour la prise de décision.

Cible : un dirigeant pressé qui doit valider l'engagement sur l'AO.
Le 1-pager présente les faits chiffrés clés + les points business + une
recommandation argumentée + les risques principaux.
"""

from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from pydantic import BaseModel, Field

from .content_loader import load_repair_profile
from .dpgf_engine import DPGFGeneree
from .extractor import AnalyseAO
from .llm import LLMClient
from .mt_engine import MemoireTechniqueGenere


class CritereJugement(BaseModel):
    libelle: str = Field(description="Libellé du critère (ex : Valeur technique).")
    ponderation: str = Field(description="Pondération (ex : '60 %' ou 'note sur 30').")


class SyntheseDirection(BaseModel):
    titre_operation: str = Field(description="Objet du marché en 1 phrase.")
    pouvoir_adjudicateur: str
    type_marche: str
    montant_max_he: str
    duree: str
    date_remise: str

    criteres_jugement: list[CritereJugement] = Field(
        description="Critères de jugement principaux avec leur pondération."
    )

    points_attention_business: list[str] = Field(
        description="3 à 5 points d'attention business spécifiques pour la direction."
    )

    atouts_candidat: list[str] = Field(
        description="3 à 5 atouts SPÉCIFIQUES du candidat sur cet AO précis, qui croisent "
        "les exigences extraites avec le profil entreprise. Pas de générique creux."
    )

    risques_principaux: list[str] = Field(
        description="2 à 4 risques ou points de vigilance pour la direction (charges, "
        "exigences inhabituelles, contraintes de planning, références à étoffer)."
    )

    montant_dqe_estime_he: str = Field(
        description="Montant DQE estimé en euros HT (chaîne formatée : '181 200 € HT')."
    )

    recommandation_go_nogo: str = Field(
        description="Recommandation argumentée en 2-3 phrases. Conclure clairement par "
        "'Recommandation : GO' ou 'Recommandation : GO sous conditions [...]' ou "
        "'Recommandation : NO-GO'."
    )


_SYSTEM_PROMPT = """Tu produis une **synthèse direction** d'un AO public pour un BET structures \
spécialisé en pathologie / diagnostic / confortement, à destination d'un dirigeant qui doit \
arbitrer go / no-go en quelques minutes.

Le ton est :
- **Factuel** — chiffres, dates, ponds, chiffrage. Pas de marketing.
- **Synthétique** — chaque bullet point fait 1 ligne, pas de prose étalée.
- **Honnête** — les risques sont explicites, pas dissimulés.

Pour les **atouts du candidat**, croise impérativement le profil entreprise (transmis en \
contexte) avec les spécificités de l'AO (extraites de l'analyse). N'écris PAS d'atouts \
génériques creux ("expertise reconnue", "équipe motivée") — n'écris que des atouts \
SPÉCIFIQUES qui correspondent à des exigences précises de cet AO.

Pour la **recommandation**, base-toi sur :
- Adéquation cœur de métier / objet du marché
- Capacité technique vs exigences
- Niveau de concurrence présumé
- Charge prévisionnelle vs capacité
- Montant et durée du marché

Conclus toujours par une formule explicite "Recommandation : GO" / "GO sous conditions" / \
"NO-GO".
"""


def generer_synthese(
    client: LLMClient,
    analyse: AnalyseAO,
    mt: MemoireTechniqueGenere | None,
    dpgf: DPGFGeneree | None,
) -> SyntheseDirection:
    """Génère la synthèse direction.

    mt et dpgf peuvent être None : dans ce cas la synthèse se fait sans visibilité
    sur le contenu déjà rédigé (utile pour un go/no-go en amont du travail de réponse).
    """
    repair_profile = load_repair_profile()

    contexte_parts = [
        "## Analyse de l'AO\n",
        "```json",
        analyse.model_dump_json(indent=2),
        "```",
    ]

    if mt:
        contexte_parts.append("\n## Sections du mémoire technique pré-rédigé\n")
        for s in mt.sections:
            contexte_parts.append(f"### {s.titre}")
            if s.sous_titre:
                contexte_parts.append(f"_{s.sous_titre}_")
            contexte_parts.append(s.contenu_md[:600] + ("…" if len(s.contenu_md) > 600 else ""))

    if dpgf:
        contexte_parts.append(
            f"\n## Programme indicatif chiffré\n\n"
            f"{dpgf.description_programme}\n\n"
            f"**Montant DQE estimé : {dpgf.montant_dqe_he:,.0f} € HT**".replace(",", " ")
        )

    if repair_profile:
        contexte_parts.append("\n## Profil entreprise\n")
        contexte_parts.append(repair_profile)

    user_prompt = (
        "\n".join(contexte_parts)
        + "\n\n---\n\nProduis la SyntheseDirection complète selon le schéma fourni."
    )

    return client.extract_structured(
        system_prompt=_SYSTEM_PROMPT,
        dce_context=None,
        user_message=user_prompt,
        schema=SyntheseDirection,
        max_tokens=4000,
        temperature=0.2,
        thinking_budget=1024,
    )


# ---------------------------------------------------------------------------
# Export DOCX 1-pager — layout exécutif en colonnes colorées
# ---------------------------------------------------------------------------

COULEUR_PRIMAIRE = RGBColor(0x1A, 0x3D, 0x6E)        # bleu marine
COULEUR_PRIMAIRE_HEX = "1A3D6E"
COULEUR_VERT = RGBColor(0x10, 0x6B, 0x3A)            # GO
COULEUR_VERT_HEX = "106B3A"
COULEUR_VERT_FOND = "E8F5EE"                         # fond Atouts
COULEUR_ORANGE = RGBColor(0xB5, 0x47, 0x08)          # vigilance
COULEUR_ORANGE_HEX = "B54708"
COULEUR_ORANGE_FOND = "FEF4E6"                       # fond Risques
COULEUR_ROUGE_HEX = "B91C1C"
COULEUR_GRIS_FOND = "F4F6FA"                         # cartes neutres
COULEUR_BLEU_FOND = "E8EDF5"                         # fond bandeau identité
COULEUR_TEXTE = RGBColor(0x22, 0x22, 0x22)


def _set_cell_shading(cell, hex_color: str) -> None:
    """Couleur de fond d'une cellule de table."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _remove_cell_borders(cell) -> None:
    """Enlève les bordures d'une cellule (sauf style explicite)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "nil")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _set_cell_left_border(cell, hex_color: str, sz: str = "16") -> None:
    """Ajoute une bordure gauche colorée (effet "card" avec accent à gauche)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    left = tcBorders.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        tcBorders.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), sz)
    left.set(qn("w:space"), "0")
    left.set(qn("w:color"), hex_color)


def _badge_recommandation(reco: str) -> tuple[str, str, str]:
    """Renvoie (label, couleur_fond_hex, couleur_texte_hex) selon la reco."""
    upper = reco.upper()
    if "NO-GO" in upper or "NO GO" in upper:
        return ("NO-GO", COULEUR_ROUGE_HEX, "FFFFFF")
    if "SOUS CONDITIONS" in upper or "CONDITIONN" in upper:
        return ("GO sous conditions", COULEUR_ORANGE_HEX, "FFFFFF")
    return ("GO", COULEUR_VERT_HEX, "FFFFFF")


def _set_table_no_borders(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            _remove_cell_borders(cell)


def _add_run(paragraph, text: str, *, bold=False, italic=False, size=10, color=None,
             family="Calibri") -> None:
    run = paragraph.add_run(text)
    run.font.name = family
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _add_card_title(cell, titre: str, color=COULEUR_PRIMAIRE) -> None:
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.text = ""
    _add_run(p, titre.upper(), bold=True, size=10, color=color)
    p.paragraph_format.space_after = Pt(4)


def _add_kv_to_cell(cell, cle: str, valeur: str) -> None:
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    _add_run(p, f"{cle} : ", bold=True, size=9)
    _add_run(p, valeur, size=9)


def _add_bullet_to_cell(cell, texte: str) -> None:
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.3)
    _add_run(p, "▪  ", size=9, color=COULEUR_PRIMAIRE)
    # Gestion **bold**
    pattern = re.compile(r"\*\*(.+?)\*\*")
    pos = 0
    for m in pattern.finditer(texte):
        if m.start() > pos:
            _add_run(p, texte[pos : m.start()], size=9)
        _add_run(p, m.group(1), bold=True, size=9)
        pos = m.end()
    if pos < len(texte):
        _add_run(p, texte[pos:], size=9)


def exporter_synthese_docx(
    synthese: SyntheseDirection, output_path: Path, candidat: str = ""
) -> Path:
    """Génère le 1-pager DOCX exécutif et retourne son chemin."""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.font.color.rgb = COULEUR_TEXTE

    # ============================================================
    # 1. Bandeau de titre (table 1 colonne, fond bleu marine)
    # ============================================================
    titre_table = doc.add_table(rows=1, cols=1)
    titre_table.autofit = False
    cell = titre_table.cell(0, 0)
    _set_cell_shading(cell, COULEUR_PRIMAIRE_HEX)
    _remove_cell_borders(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    _add_run(p, "SYNTHÈSE DIRECTION  —  Go / No-go", bold=True, size=18, color=RGBColor(0xFF, 0xFF, 0xFF))
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    sub = (
        f"{candidat}  ·  " if candidat else ""
    ) + f"Analyse du {date.today().strftime('%d/%m/%Y')}"
    _add_run(p2, sub, size=9, italic=True, color=RGBColor(0xCC, 0xD9, 0xE8))

    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_after = Pt(6)

    # ============================================================
    # 2. Bandeau Identité + Recommandation (2 cellules)
    # ============================================================
    haut = doc.add_table(rows=1, cols=2)
    haut.autofit = False
    haut.columns[0].width = Cm(11)
    haut.columns[1].width = Cm(7)
    _set_table_no_borders(haut)

    cell_id = haut.cell(0, 0)
    _set_cell_shading(cell_id, COULEUR_BLEU_FOND)
    _set_cell_left_border(cell_id, COULEUR_PRIMAIRE_HEX, sz="24")
    cell_id.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    _add_card_title(cell_id, "Identité du marché")
    _add_kv_to_cell(cell_id, "Objet", synthese.titre_operation)
    _add_kv_to_cell(cell_id, "Pouvoir adjudicateur", synthese.pouvoir_adjudicateur)
    _add_kv_to_cell(cell_id, "Type de marché", synthese.type_marche)
    _add_kv_to_cell(cell_id, "Date limite de remise", synthese.date_remise)
    _add_kv_to_cell(cell_id, "Durée", synthese.duree)

    cell_reco = haut.cell(0, 1)
    label, fond_hex, txt_hex = _badge_recommandation(synthese.recommandation_go_nogo)
    _set_cell_shading(cell_reco, fond_hex)
    _remove_cell_borders(cell_reco)
    cell_reco.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell_reco.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    _add_run(
        p, "RECOMMANDATION", bold=True, size=10,
        color=RGBColor(int(txt_hex[0:2], 16), int(txt_hex[2:4], 16), int(txt_hex[4:6], 16)),
    )
    p2 = cell_reco.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(8)
    _add_run(
        p2, label, bold=True, size=22,
        color=RGBColor(int(txt_hex[0:2], 16), int(txt_hex[2:4], 16), int(txt_hex[4:6], 16)),
    )

    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_after = Pt(4)

    # ============================================================
    # 3. Bandeau chiffres clés (3 colonnes)
    # ============================================================
    stats = doc.add_table(rows=1, cols=3)
    stats.autofit = False
    for col in stats.columns:
        col.width = Cm(6)
    _set_table_no_borders(stats)

    libs = [
        ("Montant maximum HT", synthese.montant_max_he),
        ("DQE estimé HT", synthese.montant_dqe_estime_he),
        ("Date limite", synthese.date_remise),
    ]
    for i, (label_s, value) in enumerate(libs):
        c = stats.cell(0, i)
        _set_cell_shading(c, COULEUR_GRIS_FOND)
        _remove_cell_borders(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(0)
        _add_run(p, label_s.upper(), size=8, color=RGBColor(0x66, 0x66, 0x66), bold=True)
        p2 = c.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(4)
        _add_run(p2, value, size=12, bold=True, color=COULEUR_PRIMAIRE)

    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_after = Pt(4)

    # ============================================================
    # 4. Atouts (vert) || Risques (orange)
    # ============================================================
    atouts_risques = doc.add_table(rows=1, cols=2)
    atouts_risques.autofit = False
    atouts_risques.columns[0].width = Cm(9)
    atouts_risques.columns[1].width = Cm(9)
    _set_table_no_borders(atouts_risques)

    cell_atouts = atouts_risques.cell(0, 0)
    _set_cell_shading(cell_atouts, COULEUR_VERT_FOND)
    _set_cell_left_border(cell_atouts, COULEUR_VERT_HEX, sz="24")
    _add_card_title(cell_atouts, "✓  Atouts du candidat", color=COULEUR_VERT)
    for atout in synthese.atouts_candidat:
        _add_bullet_to_cell(cell_atouts, atout)

    cell_risques = atouts_risques.cell(0, 1)
    _set_cell_shading(cell_risques, COULEUR_ORANGE_FOND)
    _set_cell_left_border(cell_risques, COULEUR_ORANGE_HEX, sz="24")
    _add_card_title(cell_risques, "⚠  Risques et vigilance", color=COULEUR_ORANGE)
    for risque in synthese.risques_principaux:
        _add_bullet_to_cell(cell_risques, risque)

    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_after = Pt(4)

    # ============================================================
    # 5. Critères de jugement || Points d'attention
    # ============================================================
    bas = doc.add_table(rows=1, cols=2)
    bas.autofit = False
    bas.columns[0].width = Cm(9)
    bas.columns[1].width = Cm(9)
    _set_table_no_borders(bas)

    cell_crit = bas.cell(0, 0)
    _set_cell_shading(cell_crit, COULEUR_GRIS_FOND)
    _set_cell_left_border(cell_crit, COULEUR_PRIMAIRE_HEX, sz="20")
    _add_card_title(cell_crit, "Critères de jugement")
    for c in synthese.criteres_jugement:
        _add_bullet_to_cell(cell_crit, f"**{c.libelle}** — {c.ponderation}")

    cell_att = bas.cell(0, 1)
    _set_cell_shading(cell_att, COULEUR_GRIS_FOND)
    _set_cell_left_border(cell_att, COULEUR_PRIMAIRE_HEX, sz="20")
    _add_card_title(cell_att, "Points d'attention business")
    for p_att in synthese.points_attention_business:
        _add_bullet_to_cell(cell_att, p_att)

    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_after = Pt(6)

    # ============================================================
    # 6. Argumentaire de la recommandation
    # ============================================================
    p = doc.add_paragraph()
    _add_run(p, "ARGUMENTAIRE DE LA RECOMMANDATION", bold=True, size=10, color=COULEUR_PRIMAIRE)
    p.paragraph_format.space_after = Pt(2)

    p_argum = doc.add_paragraph()
    p_argum.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_argum.paragraph_format.space_after = Pt(0)
    _add_run(p_argum, synthese.recommandation_go_nogo, size=10)

    doc.save(str(output_path))
    return output_path
