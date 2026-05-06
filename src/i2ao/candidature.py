"""Pack candidature complet : lettre de présentation + ZIP des livrables.

Le pack ZIP contient les pièces que le candidat envoie effectivement au pouvoir
adjudicateur :
  - lettre-presentation.docx
  - memoire-technique.docx
  - dpgf.xlsx
  - synthese-direction.docx (interne, optionnelle)
"""

from __future__ import annotations

import zipfile
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from pydantic import BaseModel, Field

from .affaires import Affaire
from .content_loader import load_repair_profile
from .extractor import AnalyseAO
from .llm import LLMClient


# ---------------------------------------------------------------------------
# Schéma de la lettre
# ---------------------------------------------------------------------------


class LettrePresentation(BaseModel):
    en_tete_candidat: str = Field(
        description="Bloc en-tête de l'expéditeur : raison sociale, adresse, contact. "
        "Format multi-lignes (\\n entre lignes)."
    )
    destinataire_bloc: str = Field(
        description="Bloc adresse du destinataire : représentant + collectivité + adresse. "
        "Format multi-lignes."
    )
    lieu_date: str = Field(
        description="Ex : 'Grenoble, le 5 mai 2026'."
    )
    objet: str = Field(
        description="Objet précis de la lettre : 'Objet : Candidature au marché [référence] - "
        "[objet du marché]'."
    )
    introduction: str = Field(
        description="1 paragraphe (3-5 phrases) : référence à la consultation, raison "
        "d'intérêt du candidat, confirmation d'engagement à respecter le cahier des charges."
    )
    presentation_candidat: str = Field(
        description="1 paragraphe (3-5 phrases) : présentation synthétique du candidat, "
        "son positionnement, ses moyens. Pas de marketing creux."
    )
    atouts_specifiques: str = Field(
        description="1 paragraphe (3-5 phrases) : atouts SPÉCIFIQUES du candidat sur cet AO "
        "précis (croisement profil × exigences DCE). Pas de générique."
    )
    pieces_jointes: list[str] = Field(
        description="Liste exhaustive et numérotée des pièces remises avec la candidature : "
        "DC1, DC2, attestations, mémoire technique, planning, BPU/DPGF, RIB, etc. Pas de "
        "sous-listes."
    )
    formule_politesse: str = Field(
        description="Formule de politesse classique pour marché public, 1 phrase."
    )
    signataire_nom: str = Field(description="Nom du signataire.")
    signataire_qualite: str = Field(description="Qualité du signataire (ex : 'Gérant').")


_SYSTEM_PROMPT = """Tu rédiges une **lettre de présentation** pour accompagner une réponse à un \
appel d'offres public français, dans le créneau bureau d'études structures pathologie / \
diagnostic / confortement.

Le ton est :
- **Sobre et professionnel** — pas de marketing creux, pas de ton commercial.
- **Précis** — on cite l'AO par sa référence, on liste les pièces concrètes.
- **Concret** — les atouts sont sourcés sur le profil candidat ET les exigences du DCE, \
pas inventés.

**Règles :**
1. Trois paragraphes de corps (introduction / présentation / atouts spécifiques) — pas plus.
2. Les atouts spécifiques doivent être **du concret** : RGA, site occupé, qualifications, \
filiale travaux, références. Pas "expertise reconnue".
3. La liste des pièces jointes est exhaustive ET réaliste (DC1, DC2, attestations, MT, \
planning, BPU/DPGF, RIB, RC le cas échéant).
4. Le destinataire est correctement identifié depuis l'analyse AO (Monsieur le Maire / \
Madame la Présidente / Monsieur le Directeur Général selon le pouvoir adjudicateur).
5. Date du jour à utiliser dans `lieu_date`.
"""


def generer_lettre(
    client: LLMClient, analyse: AnalyseAO, candidat: str = "Repair Ingénierie"
) -> LettrePresentation:
    """Génère le brouillon de lettre via Gemini."""
    profile = load_repair_profile()

    user_prompt = f"""## Analyse de l'AO

```json
{analyse.model_dump_json(indent=2)}
```

## Profil de l'entreprise candidate

{profile if profile else f"Candidat : {candidat} (profil détaillé non fourni)"}

## Date à indiquer

{date.today().strftime('%d %B %Y').replace('January', 'janvier').replace('February', 'février').replace('March', 'mars').replace('April', 'avril').replace('May', 'mai').replace('June', 'juin').replace('July', 'juillet').replace('August', 'août').replace('September', 'septembre').replace('October', 'octobre').replace('November', 'novembre').replace('December', 'décembre')}

---

Produis la LettrePresentation complète selon le schéma fourni.
"""

    return client.extract_structured(
        system_prompt=_SYSTEM_PROMPT,
        dce_context=None,
        user_message=user_prompt,
        schema=LettrePresentation,
        max_tokens=4000,
        temperature=0.3,
        thinking_budget=512,
    )


# ---------------------------------------------------------------------------
# Export DOCX de la lettre
# ---------------------------------------------------------------------------

COULEUR_PRIMAIRE = RGBColor(0x1A, 0x3D, 0x6E)
COULEUR_TEXTE = RGBColor(0x22, 0x22, 0x22)


def exporter_lettre_docx(lettre: LettrePresentation, output_path: Path) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.font.color.rgb = COULEUR_TEXTE

    # En-tête candidat (gauche)
    for ligne in lettre.en_tete_candidat.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(ligne)
        run.font.size = Pt(10)
        if ligne.strip() and ligne == lettre.en_tete_candidat.split("\n")[0]:
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = COULEUR_PRIMAIRE

    for _ in range(2):
        doc.add_paragraph()

    # Destinataire (à droite)
    for ligne in lettre.destinataire_bloc.split("\n"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(ligne)
        run.font.size = Pt(10)

    for _ in range(2):
        doc.add_paragraph()

    # Lieu et date (à droite)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(lettre.lieu_date)
    run.italic = True
    run.font.size = Pt(10)

    doc.add_paragraph()

    # Objet (gras)
    p = doc.add_paragraph()
    run = p.add_run(lettre.objet)
    run.bold = True
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(18)

    # Madame, Monsieur,
    p = doc.add_paragraph()
    p.add_run("Madame, Monsieur,").italic = True
    p.paragraph_format.space_after = Pt(12)

    # Corps : 3 paragraphes
    for paragraphe in (
        lettre.introduction,
        lettre.presentation_candidat,
        lettre.atouts_specifiques,
    ):
        p = doc.add_paragraph(paragraphe)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(0.8)
        p.paragraph_format.space_after = Pt(10)

    # Pièces jointes
    p = doc.add_paragraph("Pièces jointes à la présente candidature :")
    p.paragraph_format.space_after = Pt(6)
    for piece in lettre.pieces_jointes:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(piece).font.size = Pt(10)

    doc.add_paragraph()

    # Formule de politesse
    p = doc.add_paragraph(lettre.formule_politesse)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(36)

    # Signature (à droite)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(lettre.signataire_nom)
    run.bold = True
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(lettre.signataire_qualite)
    run.italic = True
    run.font.size = Pt(10)

    doc.save(str(output_path))
    return output_path


# ---------------------------------------------------------------------------
# Pack ZIP
# ---------------------------------------------------------------------------


def creer_pack_zip(affaire: Affaire, candidat: str = "Repair Ingenierie") -> Path:
    """Assemble en ZIP les pièces de la candidature.

    Le ZIP contient :
      - 01-lettre-presentation.docx (si générée)
      - 02-memoire-technique.docx
      - 03-dpgf.xlsx
      - 04-synthese-direction.docx (interne — informative)
    """
    pack_path = affaire.dossier / f"pack-candidature-{affaire.slug}.zip"
    cand_safe = candidat.replace(" ", "-").replace("/", "-")

    with zipfile.ZipFile(pack_path, "w", zipfile.ZIP_DEFLATED) as zf:
        if affaire.lettre_docx_path.exists():
            zf.write(
                affaire.lettre_docx_path,
                arcname=f"01-Lettre-presentation_{cand_safe}.docx",
            )
        if affaire.mt_docx_path.exists():
            zf.write(
                affaire.mt_docx_path,
                arcname=f"02-Memoire-technique_{cand_safe}.docx",
            )
        if affaire.dpgf_xlsx_path.exists():
            zf.write(
                affaire.dpgf_xlsx_path,
                arcname=f"03-DPGF_{cand_safe}.xlsx",
            )
        if affaire.synthese_docx_path.exists():
            zf.write(
                affaire.synthese_docx_path,
                arcname=f"99-Synthese-direction_INTERNE_{cand_safe}.docx",
            )

    return pack_path
