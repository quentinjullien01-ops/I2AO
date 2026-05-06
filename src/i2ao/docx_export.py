"""Export DOCX d'un mémoire technique avec mise en page sobre type marché public.

Structure :
  - Page de couverture (titre, marché, candidat, date)
  - Page de sommaire (numérotée, titres des sections)
  - Sections numérotées avec saut de page entre chacune
  - En-tête : référence marché + candidat
  - Pied de page : page X / total
"""

from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .mt_engine import MemoireTechniqueGenere, SectionMT

COULEUR_TITRE = RGBColor(0x1A, 0x3D, 0x6E)
COULEUR_TEXTE = RGBColor(0x22, 0x22, 0x22)
COULEUR_GRIS = RGBColor(0x66, 0x66, 0x66)
COULEUR_LIGNE = RGBColor(0xD0, 0xD7, 0xE3)


def exporter_mt_docx(mt: MemoireTechniqueGenere, output_path: Path) -> Path:
    """Génère le DOCX du mémoire technique et retourne son chemin."""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _configurer_styles(doc)
    _configurer_mise_en_page(doc)
    _ajouter_header_footer(doc, mt)

    _ajouter_couverture(doc, mt)
    doc.add_page_break()
    _ajouter_sommaire(doc, mt)
    doc.add_page_break()

    for index, section in enumerate(mt.sections, start=1):
        _ajouter_section(doc, section, index)
        if index < len(mt.sections):
            doc.add_page_break()

    doc.save(str(output_path))
    return output_path


# ---------------------------------------------------------------------------
# Configuration globale
# ---------------------------------------------------------------------------


def _configurer_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.font.color.rgb = COULEUR_TEXTE
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15


def _configurer_mise_en_page(doc: Document) -> None:
    for s in doc.sections:
        s.top_margin = Cm(2.2)
        s.bottom_margin = Cm(2.2)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)
        s.header_distance = Cm(1.0)
        s.footer_distance = Cm(1.0)


def _ajouter_header_footer(doc: Document, mt: MemoireTechniqueGenere) -> None:
    section = doc.sections[0]
    # Page de garde sans header/footer
    section.different_first_page_header_footer = True

    # Header (sauf 1ère page)
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"{mt.titre_marche}  ·  {mt.candidat}")
    run.font.name = "Times New Roman"
    run.font.size = Pt(8.5)
    run.font.color.rgb = COULEUR_GRIS
    run.italic = True

    # Footer (sauf 1ère page) : Page X / Y
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Page ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)
    run.font.color.rgb = COULEUR_GRIS
    _insert_field(p, "PAGE")
    run = p.add_run(" / ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)
    run.font.color.rgb = COULEUR_GRIS
    _insert_field(p, "NUMPAGES")


def _insert_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)
    run.font.color.rgb = COULEUR_GRIS
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = f" {field_code} "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


# ---------------------------------------------------------------------------
# Page de couverture
# ---------------------------------------------------------------------------


def _ajouter_couverture(doc: Document, mt: MemoireTechniqueGenere) -> None:
    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MÉMOIRE TECHNIQUE")
    run.font.name = "Times New Roman"
    run.font.size = Pt(30)
    run.bold = True
    run.font.color.rgb = COULEUR_TITRE

    _ajouter_separateur(doc, largeur_cm=8)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(mt.titre_marche)
    run.font.name = "Times New Roman"
    run.font.size = Pt(15)
    run.italic = True
    run.font.color.rgb = COULEUR_TEXTE

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Pouvoir adjudicateur")
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = COULEUR_GRIS
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(mt.pouvoir_adjudicateur)
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    run.font.color.rgb = COULEUR_TEXTE
    p.paragraph_format.space_after = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Candidat")
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = COULEUR_GRIS
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(mt.candidat)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = COULEUR_TITRE

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(date.today().strftime("%d %B %Y"))
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = COULEUR_GRIS


def _ajouter_separateur(doc: Document, largeur_cm: float = 6.0) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1A3D6E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    run = p.add_run(" ")
    run.font.size = Pt(1)


# ---------------------------------------------------------------------------
# Page de sommaire
# ---------------------------------------------------------------------------


def _ajouter_sommaire(doc: Document, mt: MemoireTechniqueGenere) -> None:
    p = doc.add_paragraph()
    run = p.add_run("SOMMAIRE")
    run.font.name = "Times New Roman"
    run.font.size = Pt(20)
    run.bold = True
    run.font.color.rgb = COULEUR_TITRE
    p.paragraph_format.space_after = Pt(20)

    for index, section in enumerate(mt.sections, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(f"{index}.  ")
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = COULEUR_TITRE
        run = p.add_run(section.titre)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.font.color.rgb = COULEUR_TEXTE
        if section.sous_titre:
            run = p.add_run(f"  —  {section.sous_titre}")
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            run.italic = True
            run.font.color.rgb = COULEUR_GRIS


# ---------------------------------------------------------------------------
# Section
# ---------------------------------------------------------------------------


def _ajouter_section(doc: Document, section: SectionMT, index: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{index}.")
    run.font.name = "Times New Roman"
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = COULEUR_TITRE
    run = p.add_run(f"  {section.titre}")
    run.font.name = "Times New Roman"
    run.font.size = Pt(20)
    run.bold = True
    run.font.color.rgb = COULEUR_TITRE

    if section.sous_titre:
        p = doc.add_paragraph()
        run = p.add_run(section.sous_titre)
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)
        run.italic = True
        run.font.color.rgb = COULEUR_GRIS
        p.paragraph_format.space_after = Pt(6)

    _ajouter_separateur(doc, largeur_cm=15)
    doc.add_paragraph()

    _rendre_contenu_md(doc, section.contenu_md)


# ---------------------------------------------------------------------------
# Rendu Markdown -> Word
# ---------------------------------------------------------------------------

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ITAL_RE = re.compile(r"(?<!\*)\*([^*]+?)\*(?!\*)")


def _rendre_contenu_md(doc: Document, md: str) -> None:
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            i += 1
            continue

        if line.lstrip().startswith(("- ", "* ")):
            while i < len(lines) and lines[i].lstrip().startswith(("- ", "* ")):
                content = lines[i].lstrip()[2:].strip()
                _ajouter_puce(doc, content, ordered=False)
                i += 1
            continue

        if re.match(r"^\d+\.\s", line.lstrip()):
            while i < len(lines) and re.match(r"^\d+\.\s", lines[i].lstrip()):
                content = re.sub(r"^\d+\.\s+", "", lines[i].lstrip()).strip()
                _ajouter_puce(doc, content, ordered=True)
                i += 1
            continue

        para_lines: list[str] = []
        while i < len(lines):
            ln = lines[i].rstrip()
            if not ln.strip():
                break
            if ln.lstrip().startswith(("- ", "* ")) or re.match(r"^\d+\.\s", ln.lstrip()):
                break
            para_lines.append(ln)
            i += 1
        text = " ".join(p.strip() for p in para_lines)
        _ajouter_paragraphe_inline(doc, text)


def _ajouter_paragraphe_inline(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.first_line_indent = Cm(0)
    _appliquer_runs_inline(p, text)


def _ajouter_puce(doc: Document, text: str, ordered: bool) -> None:
    style = "List Number" if ordered else "List Bullet"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(3)
    _appliquer_runs_inline(p, text)


def _appliquer_runs_inline(paragraph, text: str) -> None:
    pos = 0
    bold_matches = list(_BOLD_RE.finditer(text))
    ital_matches = list(_ITAL_RE.finditer(text))
    events: list[tuple[int, int, str, str]] = []
    for m in bold_matches:
        events.append((m.start(), m.end(), "bold", m.group(1)))
    for m in ital_matches:
        if not _is_inside_bold(m.start(), bold_matches):
            events.append((m.start(), m.end(), "italic", m.group(1)))
    events.sort(key=lambda e: e[0])

    for start, end, kind, content in events:
        if start > pos:
            run = paragraph.add_run(text[pos:start])
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
        run = paragraph.add_run(content)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        if kind == "bold":
            run.bold = True
        elif kind == "italic":
            run.italic = True
        pos = end
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)


def _is_inside_bold(pos: int, bold_matches) -> bool:
    return any(m.start() < pos < m.end() for m in bold_matches)
